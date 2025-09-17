import streamlit as st
from pdf2image import convert_from_path
from PIL import Image
import fitz  # PyMuPDF
import pandas as pd
import openpyxl
import matplotlib.pyplot as plt
import io
from docx import Document
import tempfile
import os

def convert_word_to_images(docx_file, page_nums, img_format):
    # This function splits docx into pages (by section breaks) and renders as images
    # For simplicity, we treat each page as a section (not 100% accurate)
    doc = Document(docx_file)
    sections = []
    section = []
    for para in doc.paragraphs:
        section.append(para.text)
        if para.text.strip() == "":  # crude page break (could be improved)
            sections.append("\n".join(section))
            section = []
    if section:
        sections.append("\n".join(section))
    images = []
    for i in page_nums:
        buf = io.BytesIO()
        plt.figure(figsize=(8.5, 11))
        plt.text(0.1, 0.9, sections[i-1], wrap=True, fontsize=12)
        plt.axis('off')
        plt.savefig(buf, format=img_format, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        img = Image.open(buf)
        images.append(img)
    return images

def convert_pdf_to_images(pdf_file, page_nums, img_format):
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
        tmp.write(pdf_file.read())
        tmp_path = tmp.name
    images = []
    all_imgs = convert_from_path(tmp_path)
    os.remove(tmp_path)
    for i in page_nums:
        img = all_imgs[i-1]
        buf = io.BytesIO()
        img.save(buf, format=img_format)
        buf.seek(0)
        images.append(Image.open(buf))
    return images

def convert_excel_to_images(excel_file, sheets, ranges, img_format):
    wb = openpyxl.load_workbook(excel_file, data_only=True)
    images = []
    for sheet_name, cell_range in zip(sheets, ranges):
        ws = wb[sheet_name]
        df = pd.DataFrame(ws.values)
        # If cell_range is specified, slice df
        if cell_range:
            start_cell, end_cell = cell_range.split(':')
            start_row = openpyxl.utils.cell.row_index_from_string(start_cell)
            start_col = openpyxl.utils.cell.column_index_from_string(start_cell)
            end_row = openpyxl.utils.cell.row_index_from_string(end_cell)
            end_col = openpyxl.utils.cell.column_index_from_string(end_cell)
            df = df.iloc[start_row-1:end_row, start_col-1:end_col]
        fig, ax = plt.subplots(figsize=(df.shape[1]*1.2, df.shape[0]*0.5))
        ax.axis('off')
        ax.table(cellText=df.values, colLabels=df.columns, loc='center')
        buf = io.BytesIO()
        plt.savefig(buf, format=img_format, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        img = Image.open(buf)
        images.append(img)
    return images

def main():
    st.set_page_config(layout="wide")
    st.title("Chuyển đổi file sang ảnh")

    # Layout: two columns with ratio 6:4
    col1, col2 = st.columns([6, 4])

    with col1:
        st.header("Khung 1: Tải file chuyển đổi")
        uploaded_file = st.file_uploader(
            "Tải lên file (.docx, .doc, .pdf, .xls, .xlsx)",
            type=["docx", "doc", "pdf", "xls", "xlsx"]
        )

    with col2:
        st.header("Khung 2: Tùy chọn chuyển đổi")
        options = {}
        img_format = st.selectbox("Định dạng ảnh đầu ra", ["PNG", "JPG"])
        if uploaded_file:
            file_ext = uploaded_file.name.split('.')[-1].lower()
            if file_ext in ['docx', 'doc']:
                # Assume each section is a page
                doc = Document(uploaded_file)
                # crude split by blank lines for demo
                num_pages = len([para for para in doc.paragraphs if para.text.strip() == ""])
                num_pages = num_pages if num_pages > 0 else 1
                st.write(f"Số trang ước tính: {num_pages}")
                all_pages = st.checkbox("Chuyển tất cả các trang", value=True)
                if not all_pages:
                    page_nums = st.text_input("Nhập số trang muốn chuyển (vd: 1,2,3)", "1")
                    selected_pages = [int(x) for x in page_nums.split(',') if x.strip().isdigit()]
                else:
                    selected_pages = list(range(1, num_pages+1))
                options = {
                    "type": "word",
                    "pages": selected_pages,
                    "img_format": img_format
                }
            elif file_ext == 'pdf':
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                    tmp.write(uploaded_file.read())
                    tmp_path = tmp.name
                doc = fitz.open(tmp_path)
                num_pages = doc.page_count
                st.write(f"Số trang: {num_pages}")
                all_pages = st.checkbox("Chuyển tất cả các trang", value=True)
                if not all_pages:
                    page_nums = st.text_input("Nhập số trang muốn chuyển (vd: 1,2,3)", "1")
                    selected_pages = [int(x) for x in page_nums.split(',') if x.strip().isdigit()]
                else:
                    selected_pages = list(range(1, num_pages+1))
                options = {
                    "type": "pdf",
                    "pages": selected_pages,
                    "img_format": img_format
                }
                doc.close()
                os.remove(tmp_path)
            elif file_ext in ['xls', 'xlsx']:
                wb = openpyxl.load_workbook(uploaded_file, data_only=True)
                sheet_names = wb.sheetnames
                selected_sheets = st.multiselect("Chọn sheet muốn chuyển", sheet_names, default=sheet_names)
                ranges = []
                for sheet in selected_sheets:
                    cell_range = st.text_input(f"Nhập vùng dữ liệu chuyển cho sheet {sheet} (vd: A1:H20), để trống nếu muốn toàn bộ sheet", "")
                    ranges.append(cell_range if cell_range else None)
                options = {
                    "type": "excel",
                    "sheets": selected_sheets,
                    "ranges": ranges,
                    "img_format": img_format
                }

    if uploaded_file and st.button("Chuyển đổi và tải về"):
        file_ext = uploaded_file.name.split('.')[-1].lower()
        if options.get("type") == "word":
            uploaded_file.seek(0)
            images = convert_word_to_images(uploaded_file, options["pages"], options["img_format"])
        elif options.get("type") == "pdf":
            uploaded_file.seek(0)
            images = convert_pdf_to_images(uploaded_file, options["pages"], options["img_format"])
        elif options.get("type") == "excel":
            uploaded_file.seek(0)
            images = convert_excel_to_images(uploaded_file, options["sheets"], options["ranges"], options["img_format"])
        else:
            st.error("Định dạng file không hỗ trợ!")
            images = []

        if images:
            for idx, img in enumerate(images):
                buf = io.BytesIO()
                img.save(buf, format=options["img_format"])
                st.image(img, caption=f"Ảnh {idx+1}", use_column_width=True)
                st.download_button(
                    label=f"Tải về Ảnh {idx+1}",
                    data=buf.getvalue(),
                    file_name=f"output_{idx+1}.{options['img_format'].lower()}",
                    mime=f"image/{options['img_format'].lower()}"
                )
        else:
            st.warning("Không tìm thấy trang/vùng dữ liệu để chuyển đổi!")

if __name__ == "__main__":
    main()
